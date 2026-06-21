#!/usr/bin/env python3
"""
Создание .doc документации для Course Work HR
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Стили
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============== TITLE ==============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Course Work — HR Management System\n')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(13, 71, 161)
run = p.add_run('Database Design and Implementation\n\n')
run.font.size = Pt(16)
run = p.add_run('Author: Ioda Aliaksei\nGroup: JA\nDate: June 2026')
run.font.size = Pt(14)

doc.add_page_break()

# ============== 1. INTRODUCTION ==============
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'This document describes the design and implementation of a relational database '
    'for an HR (Human Resources) Management System. The system supports the full lifecycle '
    'of employee management: from candidate recruitment through hiring, training, '
    'performance evaluation, and department management.'
)

doc.add_heading('1.1 Project Goal', level=2)
doc.add_paragraph(
    'To fully design and develop the database component for an HR Management Application, '
    'including OLTP (operational) and OLAP (analytical) components, ETL processes, '
    'and reporting capabilities.'
)

doc.add_heading('1.2 Technologies Used', level=2)
items = ['MySQL / MariaDB (Relational Database)',
         'SQL (DDL, DML, Stored Procedures)',
         'Power BI (Visualization)',
         'GitHub (Version Control)']
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ============== 2. OLTP DATABASE ==============
doc.add_heading('2. OLTP Database', level=1)
doc.add_paragraph(
    'The OLTP (Online Transaction Processing) database is designed to support day-to-day '
    'HR operations. It stores information about departments, positions, employees, candidates, '
    'interviews, training programs, and recruitment campaigns.'
)

doc.add_heading('2.1 OLTP Schema (3NF)', level=2)
doc.add_paragraph(
    'The database follows Third Normal Form (3NF) principles:'
)

items = [
    '1. departments — department name, location, manager',
    '2. positions — job titles with salary ranges, linked to departments',
    '3. employees — personal info, position, salary, manager hierarchy',
    '4. candidates — job applicants, status pipeline, source channel',
    '5. interviews — evaluations for candidates with scores',
    '6. applications — many-to-many between candidates and positions',
    '7. training_programs — available courses with cost and duration',
    '8. employee_training — many-to-many enrollment tracking',
    '9. performance_reviews — employee evaluations (ratings 1-5)',
    '10. recruitment_campaigns — marketing campaigns with budgets'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.2 Tables', level=2)

tables_info = [
    ('departments (8 rows)', 'dep_id PK, name, location, manager_id FK'),
    ('positions (16 rows)', 'pos_id PK, title, dept_id FK, min/max salary'),
    ('employees (15 rows)', 'emp_id PK, name, email, phone, hire_date, pos_id FK, salary, manager_id FK, status'),
    ('candidates (10 rows)', 'cand_id PK, name, email, phone, pos_id FK, status, applied_date, source'),
    ('interviews (0 rows)', 'interview_id PK, cand_id FK, interviewer_id FK, date, type, score, result'),
    ('applications (0 rows)', 'app_id PK, cand_id FK, pos_id FK, date, status'),
    ('training_programs (8 rows)', 'prog_id PK, name, description, duration, cost, start_date'),
    ('employee_training (0 rows)', 'enroll_id PK, emp_id FK, prog_id FK, dates, status, score'),
    ('performance_reviews (0 rows)', 'review_id PK, emp_id FK, date, reviewer_id FK, rating, comments'),
    ('recruitment_campaigns (5 rows)', 'camp_id PK, name, dates, budget, channel, total_apps')
]

table = doc.add_table(rows=1, cols=2, style='Light Shading Accent 1')
hdr = table.rows[0].cells
hdr[0].text = 'Table Name (Row Count)'
hdr[1].text = 'Key Columns / Relationships'
for name, cols in tables_info:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = cols

# ============== 3. DATA LOADING ==============
doc.add_heading('3. Data Loading (CSV → OLTP)', level=1)
doc.add_paragraph(
    'Data is loaded from CSV files into the OLTP database. The CSV files contain no '
    'surrogate keys — only natural keys (department names, position titles, emails).'
)

doc.add_heading('3.1 CSV Files', level=2)
items = [
    'departments.csv — department names, locations, manager names',
    'positions.csv — position titles, department names, salary ranges',
    'employees.csv — personal data, position titles, department names, manager names',
    'candidates.csv — applicant data, position titles, sources',
    'training_programs.csv — course details',
    'recruitment_campaigns.csv — campaign info'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 Loading Script', level=2)
doc.add_paragraph(
    'Script: OLTP/02_LOAD_DATA.sql\n'
    'The script is rerunnable — it uses WHERE NOT EXISTS checks to prevent duplicate inserts. '
    'Existing records are never overwritten. Foreign key constraints are temporarily disabled '
    'during bulk loading and re-enabled afterward.'
)

# ============== 4. OLAP DATABASE ==============
doc.add_heading('4. OLAP Database (Data Warehouse)', level=1)
doc.add_paragraph(
    'The OLAP database uses a Snowflake schema designed for analytical queries. '
    'It stores aggregated data separately from the operational OLTP database.'
)

doc.add_heading('4.1 Schema Components', level=2)

items = [
    'Dimensions (7): Dim_Department, Dim_Position, Dim_Employee (SCD Type 2), '
    'Dim_Date (2020-2027), Dim_Candidate, Dim_Training, Dim_Campaign',
    'Facts (2): Fact_Employee_Snapshot (monthly headcount/salary aggregates), '
    'Fact_Recruitment (candidate pipeline metrics)',
    'SCD Type 2: Dim_Employee tracks historical salary/position changes with '
    'start_date, end_date, is_current flags',
    'Bridge Table: Bridge_Employee_Training handles many-to-many relationship '
    'between employees and training programs'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2 Analytical Questions Answered', level=2)
questions = [
    'What is the monthly salary cost per department?',
    'Which positions are hardest to fill (hire rate by position)?',
    'How effective are different recruitment channels?',
    'What is the cost per hire for each campaign?',
    'Which employees have the most training?',
    'Year-over-year salary trends by department'
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

# ============== 5. ETL PROCESS ==============
doc.add_heading('5. ETL Process (OLTP → OLAP)', level=1)
doc.add_paragraph(
    'Script: ETL/03_ETL_OLTP_to_DWH.sql\n'
    'The ETL process follows these steps:'
)

steps = [
    ('1. Identify Reference Data', 'Read dimension tables from OLTP'),
    ('2. Extract', 'Select new/changed records from OLTP tables'),
    ('3. Validate', 'Check data integrity and FK constraints'),
    ('4. Transform', 'SCD Type 2 handling — close old versions, insert new; '
     'aggregate employee data by department/position/month'),
    ('5. Load', 'Insert into DWH dimension and fact tables'),
    ('6. Bridge', 'Populate many-to-many employee-training relationships')
]
for name, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(name + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'The ETL is rerunnable — it checks for existing records before inserting '
    '(WHERE NOT EXISTS) and only updates SCD Type 2 dimensions when data changes.'
)

# ============== 6. QUERIES ==============
doc.add_heading('6. SQL Queries', level=1)

doc.add_heading('6.1 OLTP Queries (4 queries)', level=2)
oltp_queries = [
    'Q1: Employee count and avg salary by department (headcount analysis)',
    'Q2: Candidate pipeline by position (hire rate, rejection rate)',
    'Q3: Training enrollment and completion rates per employee',
    'Q4: Recruitment campaign performance (cost per application, duration)'
]
for q in oltp_queries:
    doc.add_paragraph(q, style='List Bullet')

doc.add_heading('6.2 OLAP Queries (5 queries)', level=2)
olap_queries = [
    'Q1: Monthly salary cost by department with headcount',
    'Q2: Cost per employee analysis by department',
    'Q3: Recruitment effectiveness by campaign and position (cost per hire)',
    'Q4: Employee training participation (via Bridge table)',
    'Q5: Year-over-Year salary trends'
]
for q in olap_queries:
    doc.add_paragraph(q, style='List Bullet')

# ============== 7. POWER BI REPORT ==============
doc.add_heading('7. Power BI Report', level=1)
doc.add_paragraph(
    'A Power BI report connects to the HR DWH to visualize key HR metrics:'
)

items = [
    'Title: "HR Analytics Dashboard"',
    'Slicers: Department, Date (Month/Year)',
    'Visual 1: Bar chart — Headcount by Department',
    'Visual 2: Line chart — Monthly Salary Cost Trend',
    'Visual 3: Table — Recruitment Funnel (candidates → hires by position)',
    'Visual 4: Card — Total Employees, Avg Salary, Open Positions'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ============== 8. HOW TO RUN ==============
doc.add_heading('8. How to Run', level=1)
doc.add_paragraph('Follow these steps in order:')

steps = [
    '1. Create MySQL database: mysql -u root -p < OLTP/01_CREATE_SCHEMA.sql',
    '2. Load data: mysql -u root -p < OLTP/02_LOAD_DATA.sql',
    '3. Create DWH: mysql -u root -p < OLAP/01_CREATE_DWH_SCHEMA.sql',
    '4. Run ETL: mysql -u root -p < ETL/03_ETL_OLTP_to_DWH.sql',
    '5. Run OLTP queries: mysql -u root -p < Queries/04_OLTP_QUERIES.sql',
    '6. Run OLAP queries: mysql -u root -p < Queries/05_OLAP_QUERIES.sql',
    '7. Open Power BI: connect to hr_dwh database and load the report'
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# ============== 9. FILE STRUCTURE ==============
doc.add_heading('9. File Structure', level=1)

items = [
    'OLTP/01_CREATE_SCHEMA.sql — OLTP table definitions (3NF, 10 tables)',
    'OLTP/02_LOAD_DATA.sql — Load CSV data into OLTP (rerunnable)',
    'CSV_Data/*.csv — Initial data files (6 files, no surrogate keys)',
    'OLAP/01_CREATE_DWH_SCHEMA.sql — Star/snowflake schema (7 dims, 2 facts, bridge)',
    'ETL/03_ETL_OLTP_to_DWH.sql — ETL process with SCD Type 2',
    'Queries/04_OLTP_QUERIES.sql — 4 operational queries',
    'Queries/05_OLAP_QUERIES.sql — 5 analytical queries',
    'PowerBI/HR_Analytics.pbix — Power BI report',
    'Doc/CourseWork_HR_Description.docx — this document'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# Save
doc_path = '/tmp/coursework-hr/Doc/CourseWork_HR_Description.docx'
doc.save(doc_path)
print(f'Documentation saved: {doc_path}')
